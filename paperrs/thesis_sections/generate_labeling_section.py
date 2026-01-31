#!/usr/bin/env python3
"""
Generate Labeling Strategies section for thesis in Word format
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Add title
title = doc.add_heading('3.4 Labeling Strategies', level=2)

# Introduction
intro = doc.add_paragraph(
    "To categorize website performance into three distinct classes—Good, Average, and Weak—we explored "
    "three different labeling approaches. Each method brings its own perspective to the classification problem, "
    "and comparing them helps us understand which approach produces the most meaningful categories for our "
    "machine learning models."
)

# Strategy 1
doc.add_heading('Strategy 1: Tertile-Based Division', level=3)

s1_desc = doc.add_paragraph(
    "The first approach takes a straightforward statistical route. We calculated a composite performance score "
    "for each website by averaging the normalized values of key metrics like LCP, FCP, TTI, Speed Index, and CLS. "
    "Once we had these scores, we simply divided the dataset into three equal parts. The bottom third (websites "
    "with the best scores) received the 'Good' label, the middle third got 'Average,' and the top third "
    "(worst performers) were marked as 'Weak.'"
)

s1_steps = doc.add_paragraph()
s1_steps.add_run("The process works as follows:\n").bold = True
s1_steps.add_run(
    "1. Normalize each performance metric using min-max scaling\n"
    "2. Calculate composite score by averaging the normalized values\n"
    "3. Sort all websites by their composite scores\n"
    "4. Split into three equal groups (tertiles)\n"
    "5. Assign labels: bottom tertile = Good, middle = Average, top = Weak"
)

s1_pros = doc.add_paragraph()
s1_pros.add_run("Strengths: ").bold = True
s1_pros.add_run(
    "This method is easy to implement and produces perfectly balanced classes with exactly the same number "
    "of samples in each category. It requires no assumptions about the data structure."
)

s1_cons = doc.add_paragraph()
s1_cons.add_run("Limitations: ").bold = True
s1_cons.add_run(
    "The approach treats all metrics equally, which might not reflect real-world importance. It also doesn't "
    "consider whether natural performance clusters exist in the data—it just forces three equal groups regardless "
    "of actual performance patterns."
)

# Strategy 2
doc.add_heading('Strategy 2: Weighted Composite Score', level=3)

s2_desc = doc.add_paragraph(
    "The second strategy incorporates domain expertise by assigning different importance weights to various metrics. "
    "Since Google emphasizes Core Web Vitals as critical performance indicators, we gave heavier weights to LCP, "
    "FID/INP, and CLS compared to other metrics. This creates a composite score that better reflects what actually "
    "matters for user experience."
)

s2_formula = doc.add_paragraph()
s2_formula.add_run("The weighted formula we used:\n").bold = True
s2_formula.add_run(
    "Composite Score = (0.30 × LCP) + (0.30 × FID/INP) + (0.25 × CLS) + (0.15 × other metrics)"
)

s2_thresh = doc.add_paragraph()
s2_thresh.add_run("Classification thresholds:\n").bold = True
s2_thresh.add_run(
    "• Good: Score ≥ 0.75 (meets or exceeds recommended performance)\n"
    "• Average: 0.50 ≤ Score < 0.75 (acceptable but needs improvement)\n"
    "• Weak: Score < 0.50 (requires significant optimization)"
)

s2_process = doc.add_paragraph()
s2_process.add_run("Implementation steps:\n").bold = True
s2_process.add_run(
    "1. Normalize each metric to 0-1 range\n"
    "2. Apply weights according to metric importance\n"
    "3. Calculate weighted composite score for each website\n"
    "4. Classify based on predefined thresholds aligned with Google's recommendations"
)

s2_pros = doc.add_paragraph()
s2_pros.add_run("Strengths: ").bold = True
s2_pros.add_run(
    "This method captures domain knowledge about which metrics matter most for user experience. The thresholds "
    "align with industry-standard recommendations from Google, making the labels more interpretable and actionable."
)

s2_cons = doc.add_paragraph()
s2_cons.add_run("Limitations: ").bold = True
s2_cons.add_run(
    "The weight assignments require expert judgment and might introduce bias. Different experts might choose "
    "different weights, and there's no single 'correct' weight distribution. The resulting class distribution "
    "may be imbalanced if the dataset doesn't naturally fall into these threshold ranges."
)

# Strategy 3
doc.add_heading('Strategy 3: K-means Clustering', level=3)

s3_desc = doc.add_paragraph(
    "Our third approach lets the data speak for itself through unsupervised learning. Rather than imposing "
    "predefined boundaries or weights, K-means clustering discovers natural groupings in the performance data. "
    "This method identifies three distinct clusters based on the inherent similarities and differences in the "
    "performance patterns across all metrics."
)

# Pseudocode
doc.add_heading('Algorithm Pseudocode:', level=4)

pseudo = doc.add_paragraph()
pseudo_style = pseudo.paragraph_format
pseudo_style.left_indent = Inches(0.5)
pseudo.add_run(
    "STEP 1: Data Preparation\n"
    "   Load performance metrics for all websites\n"
    "   Apply StandardScaler to normalize features (zero mean, unit variance)\n"
    "   Store normalized data in matrix X\n\n"
    
    "STEP 2: K-means Clustering\n"
    "   Initialize K-means with k=3 clusters\n"
    "   Set random_state=42 for reproducibility\n"
    "   Set n_init=10 to run algorithm 10 times with different initializations\n"
    "   Fit K-means model on normalized data X\n"
    "   Obtain cluster assignments for each website\n\n"
    
    "STEP 3: Cluster Ranking\n"
    "   FOR each cluster (0, 1, 2):\n"
    "      Extract all websites in this cluster\n"
    "      Calculate mean of all performance metrics\n"
    "      Store overall cluster mean\n"
    "   END FOR\n\n"
    
    "STEP 4: Label Assignment\n"
    "   Identify cluster with lowest mean (best performance) → label as 'Good'\n"
    "   Identify cluster with middle mean → label as 'Average'\n"
    "   Identify cluster with highest mean (worst performance) → label as 'Weak'\n\n"
    
    "STEP 5: Apply Labels\n"
    "   Map each website's cluster assignment to its corresponding label\n"
    "   Return final labeled dataset"
)

# Set monospace font for pseudocode
for run in pseudo.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

s3_explain = doc.add_paragraph(
    "The algorithm groups websites into three clusters by minimizing the within-cluster variance—essentially "
    "finding groups where websites have similar performance characteristics. After clustering, we rank the "
    "three clusters by their average performance scores to determine which represents 'Good,' 'Average,' and "
    "'Weak' performance."
)

s3_pros = doc.add_paragraph()
s3_pros.add_run("Strengths: ").bold = True
s3_pros.add_run(
    "This is a completely data-driven approach that discovers natural boundaries without human bias. It can "
    "reveal performance patterns that might not be obvious with predetermined thresholds. The method considers "
    "all features simultaneously when forming clusters, capturing complex multivariate relationships."
)

s3_cons = doc.add_paragraph()
s3_cons.add_run("Limitations: ").bold = True
s3_cons.add_run(
    "K-means is sensitive to initialization, though we control this with a fixed random seed. The algorithm "
    "assumes roughly spherical clusters, which might not match the actual data distribution. Class sizes aren't "
    "guaranteed to be balanced—one cluster might be much larger than others."
)

# Comparison and Results
doc.add_heading('Comparative Analysis', level=3)

comparison = doc.add_paragraph(
    "After implementing all three strategies, we trained identical machine learning models on each labeled "
    "dataset and compared their predictive performance. The K-means clustering approach consistently produced "
    "the best results across all evaluation metrics. This suggests that the natural groupings discovered by "
    "unsupervised learning create more meaningful and separable performance categories than artificially imposed "
    "boundaries."
)

conclusion = doc.add_paragraph(
    "The superior performance of K-means labels indicates that website performance exists in natural clusters "
    "rather than evenly distributed ranges. These clusters likely represent fundamentally different optimization "
    "states—well-optimized sites, partially optimized sites, and unoptimized sites—rather than arbitrary "
    "divisions along a continuous spectrum. For this reason, we selected the K-means labeled dataset for our "
    "final model training and evaluation."
)

# Feature Analysis Section
doc.add_page_break()
doc.add_heading('Feature Analysis', level=2)

fa_desc = doc.add_paragraph(
    "Before training our models, we examined the relationships between different performance metrics using "
    "correlation analysis. Understanding these relationships is important because highly correlated features "
    "can create redundancy in the model, potentially leading to overfitting or inflated feature importance scores."
)

fa_method = doc.add_paragraph(
    "We computed a Pearson correlation matrix across all numeric features in the dataset. This matrix reveals "
    "both positive and negative linear relationships between metrics. For instance, we expected to find strong "
    "correlations between related metrics like LCP and FCP, since both measure aspects of loading performance. "
    "Similarly, Time to Interactive and Total Blocking Time typically correlate because both reflect main thread "
    "activity."
)

fa_findings = doc.add_paragraph()
fa_findings.add_run("Key Observations:\n").bold = True
fa_findings.add_run(
    "The correlation analysis revealed several metric groups with moderate to strong relationships. Loading metrics "
    "(LCP, FCP, TTI) showed positive correlations, as did resource-based metrics (page size, number of requests). "
    "However, we didn't find extreme correlations (r > 0.95) that would warrant removing features. Each metric "
    "contributes unique information about a different aspect of performance—visual loading, interactivity, "
    "layout stability, or resource efficiency."
)

fa_decision = doc.add_paragraph(
    "Based on this analysis, we retained all 22 features for model training. While some metrics show moderate "
    "correlation, each measures a distinct aspect of web performance that the model can learn from. Modern "
    "ensemble methods like Random Forest and LightGBM handle correlated features well through their built-in "
    "feature selection mechanisms."
)

# Data Splitting Section
doc.add_heading('Data Splitting', level=2)

ds_intro = doc.add_paragraph(
    "After labeling and feature analysis, we divided the dataset into separate training and testing subsets. "
    "This split is fundamental to machine learning—we train the model on one portion of data and evaluate it "
    "on completely unseen data to assess how well it generalizes to new websites."
)

ds_method = doc.add_paragraph()
ds_method.add_run("Stratified Splitting Approach:\n").bold = True
ds_method.add_run(
    "Rather than randomly splitting the data, we used stratified sampling to ensure both subsets maintain the "
    "same proportion of Good, Average, and Weak labels. For example, if 35% of our full dataset is labeled 'Good,' "
    "then 35% of both the training and testing sets will be 'Good.' This prevents scenarios where one class might "
    "be underrepresented in the test set, which would make performance metrics unreliable."
)

ds_config = doc.add_paragraph()
ds_config.add_run("Final Split Configuration:\n").bold = True
ds_config.add_run(
    "• Training set: 80% of the data (933 website instances)\n"
    "• Testing set: 20% of the data (234 website instances)\n"
    "• Random state: 42 (for reproducibility)\n"
    "• Stratification: Applied based on performance label"
)

ds_rationale = doc.add_paragraph(
    "The 80-20 split is standard practice in machine learning. It gives the model enough data to learn complex "
    "patterns while reserving sufficient examples for reliable evaluation. With 933 training instances, the model "
    "has adequate samples from each performance category to learn distinguishing characteristics. The 234 test "
    "instances provide enough statistical power to confidently assess model accuracy and generalization capability."
)

ds_validation = doc.add_paragraph(
    "We also verified that the stratification worked correctly by checking the class distribution in both sets. "
    "The distributions matched within 1-2%, confirming that our test results would fairly represent model performance "
    "across all performance categories rather than being skewed by imbalanced test data."
)

# Save document
output_path = 'f:/client/Optimizer/optimizer/paperrs/thesis_sections/Labeling_and_Data_Preparation.docx'
doc.save(output_path)
print(f"Document created successfully: {output_path}")
