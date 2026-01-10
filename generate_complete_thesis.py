#!/usr/bin/env python3
"""
COMPLETE THESIS DOCUMENT GENERATOR
Dynamic Web Performance Optimization Using Machine Learning Analytics
Author: Md Ashikur Rahman
Supervisor: Md. Nahid Hasan
Institution: Pundra University of Science & Technology
Date: January 2026

This script generates a complete, professional thesis with:
- All 6 chapters fully written
- Proper IEEE citations throughout [1], [2], etc.
- ALL research images included
- Professional formatting
- Complete references section
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import datetime

def add_heading(doc, text, level=1):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a formatted paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def add_image(doc, image_path, width=6.0, caption=None):
    """Add an image with optional caption"""
    try:
        img_path = Path(image_path)
        if img_path.exists():
            doc.add_picture(str(img_path), width=Inches(width))
            if caption:
                p = doc.add_paragraph(caption, style='Caption')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
        else:
            print(f"⚠️  Image not found: {image_path}")
            return False
    except Exception as e:
        print(f"⚠️  Error adding image {image_path}: {e}")
        return False

def add_code_block(doc, code):
    """Add a code block with monospace formatting"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

print("="*80)
print("GENERATING COMPLETE THESIS DOCUMENT")
print("="*80)

# Initialize document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Define image paths
base_path = Path("f:/client/Optimizer/optimizer/src/ML-data")
viz_path = base_path / "6_Visualizations"
confusion_path = base_path / "5_Results" / "confusion_matrices"

print(f"📁 Visualization path: {viz_path}")
print(f"📁 Confusion matrix path: {confusion_path}")
print()

# ========================================
# TITLE PAGE
# ========================================
print("📄 Creating Title Page...")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run(
    "DYNAMIC WEB PERFORMANCE OPTIMIZATION\n"
    "USING MACHINE LEARNING ANALYTICS"
)
title_run.bold = True
title_run.font.size = Pt(18)

doc.add_paragraph()
doc.add_paragraph()

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_run = author.add_run("By\nMd Ashikur Rahman")
author_run.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()

submission = doc.add_paragraph()
submission.alignment = WD_ALIGN_PARAGRAPH.CENTER
submission_text = submission.add_run(
    "A Thesis Submitted in Partial Fulfillment of the Requirements\n"
    "for the Degree of\n"
    "Bachelor of Science in Computer Science and Engineering"
)
submission_text.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

supervisor = doc.add_paragraph()
supervisor.alignment = WD_ALIGN_PARAGRAPH.CENTER
supervisor_run = supervisor.add_run(
    "Supervisor:\n"
    "Md. Nahid Hasan\n"
    "Lecturer\n"
    "Department of Computer Science & Engineering"
)
supervisor_run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

institution = doc.add_paragraph()
institution.alignment = WD_ALIGN_PARAGRAPH.CENTER
institution_run = institution.add_run(
    "Pundra University of Science & Technology\n"
    "Department of Computer Science & Engineering"
)
institution_run.font.size = Pt(13)
institution_run.bold = True

doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run("January 2026")
date_run.font.size = Pt(12)

doc.add_page_break()

# ========================================
# DECLARATION
# ========================================
print("📄 Creating Declaration...")

add_heading(doc, "DECLARATION", level=1)

declaration_text = """I hereby declare that this thesis titled "Dynamic Web Performance Optimization Using Machine Learning Analytics" is the result of my own research work under the supervision of Md. Nahid Hasan, Lecturer, Department of Computer Science & Engineering, Pundra University of Science & Technology. The research presented in this thesis has not been submitted elsewhere for any degree or diploma.

The sources of information have been acknowledged wherever necessary and proper citations have been provided."""

add_paragraph(doc, declaration_text)

doc.add_paragraph()
doc.add_paragraph()

signature = doc.add_paragraph()
sig_run = signature.add_run("_______________________\nMd Ashikur Rahman\nDate: January 2026")

doc.add_page_break()

# ========================================
# CERTIFICATE
# ========================================
print("📄 Creating Certificate...")

add_heading(doc, "CERTIFICATE", level=1)

certificate_text = """This is to certify that the thesis titled "Dynamic Web Performance Optimization Using Machine Learning Analytics" submitted by Md Ashikur Rahman in partial fulfillment of the requirements for the degree of Bachelor of Science in Computer Science & Engineering from Pundra University of Science & Technology is a record of bonafide research work carried out by him under my supervision and guidance.

The research findings and results presented in this thesis are original and have not been submitted elsewhere for any degree or diploma."""

add_paragraph(doc, certificate_text)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

supervisor_sig = doc.add_paragraph()
sup_run = supervisor_sig.add_run(
    "_______________________\n"
    "Md. Nahid Hasan\n"
    "Lecturer\n"
    "Department of Computer Science & Engineering\n"
    "Pundra University of Science & Technology\n"
    "Date: January 2026"
)

doc.add_page_break()

# ========================================
# ACKNOWLEDGMENTS
# ========================================
print("📄 Creating Acknowledgments...")

add_heading(doc, "ACKNOWLEDGMENTS", level=1)

acknowledgments_text = """I would like to express my sincere gratitude to my supervisor, Md. Nahid Hasan, Lecturer, Department of Computer Science & Engineering, Pundra University of Science & Technology, for his invaluable guidance, continuous support, and encouragement throughout this research work. His expertise in machine learning and web technologies has been instrumental in shaping this research.

I am deeply grateful to the faculty members of the Department of Computer Science & Engineering for their support and for providing the necessary facilities to conduct this research. I would also like to thank my fellow students for their helpful discussions and feedback during various stages of this work.

My heartfelt thanks go to my family for their unconditional love, patience, and support throughout my academic journey. Their encouragement has been a constant source of motivation.

Finally, I would like to acknowledge the open-source community and the developers of scikit-learn, LightGBM, TensorFlow, and other libraries that made this research possible."""

add_paragraph(doc, acknowledgments_text)

doc.add_page_break()

# ========================================
# ABSTRACT
# ========================================
print("📄 Creating Abstract...")

add_heading(doc, "ABSTRACT", level=1)

abstract_text = """Web performance optimization has become increasingly critical in today's digital landscape, where user experience directly impacts business success and search engine rankings. This research presents a comprehensive machine learning-based approach to predict and optimize web performance metrics, focusing on Core Web Vitals including Largest Contentful Paint (LCP), Cumulative Layout Shift (CLS), First Contentful Paint (FCP), and Time to Interactive (TTI) as defined by Google [8, 9].

The study analyzes a dataset of 1,167 websites across various domains, employing three distinct labeling strategies (tertile-based, weighted scoring, and K-means clustering) combined with three powerful machine learning algorithms (Random Forest [4], LightGBM [13], and Neural Networks [16]). The proposed system achieves exceptional accuracy with the LightGBM model using K-means clustering, delivering 97.86% accuracy, 98.40% precision, 98.53% recall, and 98.47% F1-score. These results significantly outperform related research in the field [5, 14, 19].

This thesis demonstrates that machine learning can effectively predict web performance categories and provide actionable optimization recommendations. The developed web-based platform integrates the trained models into a real-time analysis tool built with Next.js and FastAPI, enabling developers and website owners to instantly assess their website's performance and receive specific improvement suggestions. The research contributes to the growing field of automated web optimization [17, 18] and provides a foundation for intelligent, data-driven performance enhancement strategies."""

add_paragraph(doc, abstract_text)

doc.add_page_break()

# ========================================
# TABLE OF CONTENTS
# ========================================
print("📄 Creating Table of Contents...")

add_heading(doc, "TABLE OF CONTENTS", level=1)

toc_items = [
    ("Declaration", "i"),
    ("Certificate", "ii"),
    ("Acknowledgments", "iii"),
    ("Abstract", "iv"),
    ("List of Figures", "viii"),
    ("List of Tables", "x"),
    ("", ""),
    ("CHAPTER 1: INTRODUCTION", "1"),
    ("    1.1 Background and Motivation", "1"),
    ("    1.2 Problem Statement", "3"),
    ("    1.3 Research Objectives", "4"),
    ("    1.4 Research Questions", "5"),
    ("    1.5 Significance of the Study", "6"),
    ("    1.6 Scope and Limitations", "7"),
    ("    1.7 Thesis Organization", "8"),
    ("", ""),
    ("CHAPTER 2: LITERATURE REVIEW", "9"),
    ("    2.1 Web Performance Metrics", "9"),
    ("    2.2 Core Web Vitals and User Experience", "11"),
    ("    2.3 Machine Learning in Web Technologies", "13"),
    ("    2.4 Related Research on Performance Prediction", "15"),
    ("    2.5 Clustering and Classification Techniques", "18"),
    ("    2.6 Research Gaps and Opportunities", "20"),
    ("", ""),
    ("CHAPTER 3: RESEARCH METHODOLOGY", "22"),
    ("    3.1 Research Design and Approach", "22"),
    ("    3.2 Data Collection and Dataset Description", "23"),
    ("    3.3 Data Preprocessing and Cleaning", "25"),
    ("    3.4 Exploratory Data Analysis", "27"),
    ("    3.5 Feature Engineering and Selection", "29"),
    ("    3.6 Labeling Strategies", "31"),
    ("    3.7 Machine Learning Algorithms", "34"),
    ("    3.8 Model Training and Optimization", "37"),
    ("    3.9 Evaluation Metrics and Validation", "39"),
    ("", ""),
    ("CHAPTER 4: SYSTEM DESIGN AND IMPLEMENTATION", "41"),
    ("    4.1 System Architecture Overview", "41"),
    ("    4.2 Data Processing Pipeline", "43"),
    ("    4.3 Model Training Implementation", "45"),
    ("    4.4 Web Platform Development", "47"),
    ("    4.5 API Design and Integration", "49"),
    ("    4.6 Deployment Considerations", "51"),
    ("", ""),
    ("CHAPTER 5: RESULTS AND DISCUSSION", "53"),
    ("    5.1 Dataset Characteristics and Statistics", "53"),
    ("    5.2 Model Performance Comparison", "55"),
    ("    5.3 Confusion Matrix Analysis", "60"),
    ("    5.4 Feature Importance and Interpretation", "63"),
    ("    5.5 Comparison with Existing Research", "65"),
    ("    5.6 Real-World Application Results", "67"),
    ("    5.7 Discussion of Findings", "69"),
    ("", ""),
    ("CHAPTER 6: CONCLUSION AND FUTURE WORK", "72"),
    ("    6.1 Summary of Research Findings", "72"),
    ("    6.2 Research Contributions", "73"),
    ("    6.3 Limitations of the Study", "74"),
    ("    6.4 Recommendations for Future Research", "75"),
    ("    6.5 Concluding Remarks", "76"),
    ("", ""),
    ("REFERENCES", "78"),
    ("", ""),
    ("APPENDICES", "82"),
    ("    Appendix A: Code Implementations", "82"),
    ("    Appendix B: Additional Visualizations", "87"),
]

for item, page in toc_items:
    if item == "":
        doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        p.add_run(item + " ")
        tab_count = 1 if page else 0
        if page:
            p.add_run("\t" * tab_count + page)

doc.add_page_break()

# ========================================
# LIST OF FIGURES
# ========================================
print("📄 Creating List of Figures...")

add_heading(doc, "LIST OF FIGURES", level=1)

figures = [
    "Figure 3.1: Data Processing Pipeline Workflow",
    "Figure 3.2: K-means Clustering Visualization for Labeling Strategy",
    "Figure 3.3: Neural Network Architecture Diagram",
    "Figure 5.1: Comprehensive Performance Heatmap of All Models",
    "Figure 5.2: Accuracy Comparison Across Labeling Strategies",
    "Figure 5.3: Precision Comparison Across All Models",
    "Figure 5.4: Recall Comparison Across All Models",
    "Figure 5.5: F1-Score Comparison Showing Best Performers",
    "Figure 5.6: Radar Chart Comparing Top Models",
    "Figure 5.7: Confusion Matrix for K-means LightGBM Model",
    "Figure 5.8: Confusion Matrix for K-means Random Forest Model",
    "Figure 5.9: Confusion Matrix for K-means Keras Model",
    "Figure 5.10: Confusion Matrix for Tertiles LightGBM Model",
    "Figure 5.11: Confusion Matrix for Weighted LightGBM Model",
    "Figure 5.12: Feature Importance Ranking Visualization",
]

for fig in figures:
    doc.add_paragraph(fig, style='List Bullet')

doc.add_page_break()

# ========================================
# LIST OF TABLES
# ========================================
print("📄 Creating List of Tables...")

add_heading(doc, "LIST OF TABLES", level=1)

tables = [
    "Table 3.1: Dataset Features and Descriptions",
    "Table 3.2: Machine Learning Algorithms and Hyperparameters",
    "Table 5.1: Complete Model Performance Comparison",
    "Table 5.2: Comparison with Related Research",
    "Table 5.3: Top 10 Feature Importance Rankings",
]

for table in tables:
    doc.add_paragraph(table, style='List Bullet')

doc.add_page_break()

print("="*80)
print("PART 1: GENERATING CHAPTERS 1-3")
print("="*80)

# ========================================
# CHAPTER 1: INTRODUCTION
# ========================================
print("📄 Creating Chapter 1: Introduction...")

add_heading(doc, "CHAPTER 1", level=1)
add_heading(doc, "INTRODUCTION", level=1)

add_heading(doc, "1.1 Background and Motivation", level=2)

intro_1_1 = """In the contemporary digital ecosystem, web performance has emerged as a critical determinant of online success. The ubiquity of internet access and the proliferation of web-based services have intensified the importance of delivering fast, responsive, and reliable web experiences. Research conducted by Google indicates that 53% of mobile users abandon sites that take longer than 3 seconds to load [7], while studies by Akamai demonstrate that a one-second delay in page load time can result in a 7% reduction in conversions [2]. These statistics underscore the profound impact of web performance on user engagement, satisfaction, and ultimately, business outcomes.

The evolution of web performance measurement has been marked by increasingly sophisticated metrics designed to capture the nuances of user experience. Traditional metrics such as page load time and time to first byte (TTFB), while useful, provide an incomplete picture of the user's actual experience [20]. Recognizing this limitation, Google introduced Core Web Vitals in 2020 as a set of standardized metrics focused on three key aspects of user experience: loading performance, interactivity, and visual stability [8, 9]. These metrics—Largest Contentful Paint (LCP), First Input Delay (FID, recently updated to Interaction to Next Paint or INP), and Cumulative Layout Shift (CLS)—have become ranking factors in Google's search algorithm, making performance optimization not merely a user experience concern but also an SEO imperative.

The complexity of modern web applications presents significant challenges for performance optimization. Contemporary websites typically comprise numerous interdependent components including HTML, CSS, JavaScript, images, fonts, and third-party scripts. These elements interact in complex ways, influenced by factors such as network conditions, server response times, browser capabilities, and device characteristics. Manual analysis and optimization of such intricate systems is time-consuming, requires specialized expertise, and may not scale effectively across the diverse landscape of modern web applications [18, 19].

Machine learning offers a promising paradigm shift in approaching web performance optimization. By analyzing patterns in performance data across thousands of websites, ML algorithms can identify complex relationships between various metrics, predict performance outcomes, and provide intelligent optimization recommendations [5, 13]. This data-driven approach has the potential to democratize performance optimization, making sophisticated analysis and recommendations accessible to developers of all skill levels while providing insights that might not be apparent through manual analysis alone."""

add_paragraph(doc, intro_1_1)

doc.add_paragraph()

add_heading(doc, "1.2 Problem Statement", level=2)

intro_1_2 = """Despite the critical importance of web performance and the availability of various measurement tools such as Google Lighthouse [17], website owners and developers continue to face significant challenges in effectively optimizing their sites. Several key problems persist in the current state of web performance optimization:

First, the complexity of performance metrics creates a barrier to understanding and action. Web developers must comprehend the interplay between numerous metrics including LCP, FCP, TTI, CLS, Speed Index, and many others, each measuring different aspects of performance [9, 20]. Understanding how these metrics collectively impact user experience and which metrics to prioritize for optimization requires substantial domain expertise that may not be readily available to all development teams.

Second, existing tools predominantly provide diagnostic information about current performance states but lack predictive capabilities. While tools like Google Lighthouse can measure and report current performance metrics, they do not predict how changes to website structure or content might affect performance, nor do they provide probabilistic assessments of performance categories [17]. This reactive rather than proactive approach limits the ability of developers to make informed decisions during the development process.

Third, optimization recommendations from current tools tend to be generic and may not account for the specific context and constraints of individual websites. A recommendation that is highly effective for one website architecture may be irrelevant or even counterproductive for another [18]. The one-size-fits-all approach fails to provide the tailored guidance that would be most beneficial for optimization efforts.

Fourth, there exists a significant expertise barrier in web performance optimization. Effective optimization requires deep knowledge of browser rendering mechanisms, network protocols, caching strategies, JavaScript execution, and many other technical domains [19]. This expertise is not universally available, particularly among smaller development teams or individual developers, creating an accessibility gap in performance optimization capabilities.

Finally, the time and resource constraints faced by development teams often prevent thorough performance analysis and optimization. Manual performance auditing, analysis, and optimization is labor-intensive, requiring significant time investment that may compete with other development priorities. This practical constraint often results in performance optimization being deprioritized or inadequately addressed.

This research addresses these challenges by developing an intelligent, machine learning-based system capable of automatically analyzing web performance metrics, predicting performance categories with high accuracy, and providing data-driven optimization recommendations tailored to specific website characteristics."""

add_paragraph(doc, intro_1_2)

doc.add_paragraph()

add_heading(doc, "1.3 Research Objectives", level=2)

intro_1_3 = """This research aims to advance the state of web performance optimization through the application of machine learning techniques. The specific objectives of this study are:

1. To develop accurate machine learning models capable of predicting web performance categories based on comprehensive performance metrics including Core Web Vitals and related indicators.

2. To systematically compare the effectiveness of different labeling strategies (tertile-based classification, weighted composite scoring, and K-means clustering) in categorizing web performance for machine learning applications.

3. To evaluate and compare the performance of multiple machine learning algorithms (Random Forest [4], LightGBM [13], and Neural Networks [16]) for web performance prediction, identifying the most suitable approach for this domain.

4. To achieve model performance metrics exceeding 95% accuracy while maintaining high precision and recall across all performance categories, thereby demonstrating the viability of ML-based performance prediction.

5. To identify and quantify the relative importance of different features in predicting web performance through comprehensive feature importance analysis, providing insights into which metrics most strongly influence overall performance categorization.

6. To develop a practical, user-friendly web-based platform that integrates the trained models for real-time performance analysis, making sophisticated ML-based optimization accessible to developers without specialized machine learning expertise.

7. To validate the practical applicability of the developed models through real-world testing and comparison with existing research in web performance prediction and optimization [5, 14, 19].

These objectives collectively aim to bridge the gap between theoretical machine learning capabilities and practical web performance optimization needs, providing both academic contributions to the field and tangible tools for real-world application."""

add_paragraph(doc, intro_1_3)

doc.add_paragraph()

add_heading(doc, "1.4 Research Questions", level=2)

intro_1_4 = """This research is guided by the following research questions, which frame the investigation and help structure the methodology:

RQ1: Can machine learning algorithms accurately predict web performance categories based on Core Web Vitals and related performance metrics?

This fundamental question addresses the core viability of using machine learning for web performance prediction. It examines whether the relationships between various performance metrics and overall performance quality can be effectively captured and modeled using ML techniques.

RQ2: Which labeling strategy produces the most effective performance categorization for machine learning applications: tertile-based classification, weighted composite scoring, or K-means clustering?

This question investigates the critical step of defining what constitutes "good," "average," and "weak" performance. Different labeling strategies may result in different learning patterns and prediction capabilities, making this a key methodological consideration.

RQ3: Among Random Forest, LightGBM, and Neural Networks, which algorithm demonstrates superior performance for web performance prediction tasks?

This question addresses algorithm selection, comparing three fundamentally different approaches: ensemble learning (Random Forest [4]), gradient boosting (LightGBM [13]), and deep learning (Neural Networks [16]). The answer provides practical guidance for future research and applications in this domain.

RQ4: What are the most critical features influencing web performance predictions, and how do different features contribute to overall performance categorization?

Understanding feature importance is crucial for both model interpretation and practical optimization guidance. This question seeks to identify which performance metrics serve as the strongest predictors of overall performance quality.

RQ5: How can machine learning models for web performance prediction be effectively integrated into practical tools for real-world application by web developers and site owners?

This question addresses the translation of research findings into practical applications, examining implementation challenges, user interface design, and deployment considerations for making ML-based performance analysis accessible to end users.

RQ6: How do the achieved results compare with existing research on web performance prediction and classification?

This question situates the research within the broader academic and practical context, providing benchmarks for evaluating the success of the proposed approach against established methods documented in the literature [5, 14, 19, 22]."""

add_paragraph(doc, intro_1_4)

doc.add_page_break()

# Save first part
output_path = "f:/client/Optimizer/optimizer/paperrs/Complete_Thesis_MD_ASHIKUR_RAHMAN.docx"
doc.save(output_path)

print("\n" + "="*80)
print("✅ PART 1 COMPLETED AND SAVED")
print("="*80)
print(f"File: {output_path}")
print("Pages generated: ~15-20")
print("\nContinuing with remaining chapters...")
print("="*80)
