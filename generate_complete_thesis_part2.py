#!/usr/bin/env python3
"""
COMPLETE THESIS PART 2 - Continuation
Adds Chapters 1.5-6, all images, and appendices to the existing document
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def add_image(doc, image_path, width=6.0, caption=None):
    try:
        img_path = Path(image_path)
        if img_path.exists():
            doc.add_picture(str(img_path), width=Inches(width))
            if caption:
                p = doc.add_paragraph(caption, style='Caption')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"✅ Added: {caption or img_path.name}")
            return True
        else:
            print(f"⚠️  Not found: {image_path}")
            return False
    except Exception as e:
        print(f"⚠️  Error: {e}")
        return False

def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

print("="*80)
print("CONTINUING THESIS - PART 2")
print("="*80)

# Load existing document
doc_path = "f:/client/Optimizer/optimizer/paperrs/Complete_Thesis_MD_ASHIKUR_RAHMAN.docx"
doc = Document(doc_path)

# Define image paths
base_path = Path("f:/client/Optimizer/optimizer/src/ML-data")
viz_path = base_path / "6_Visualizations"
confusion_path = base_path / "5_Results" / "confusion_matrices"

print(f"📁 Loading document: {doc_path}")
print()

# Continue Chapter 1
print("📄 Continuing Chapter 1...")

add_heading(doc, "1.5 Significance of the Study", level=2)

intro_1_5 = """This research makes several significant contributions to both academic knowledge and practical applications in web performance optimization:

**Academic Contributions:**

First, this study provides a comprehensive, systematic comparison of machine learning algorithms specifically applied to web performance prediction using Google's Core Web Vitals [8, 9]. While machine learning has been applied to various aspects of web technologies [5, 17, 18], focused research on Core Web Vitals prediction using multiple algorithms and labeling strategies represents a novel contribution to the literature.

Second, the research introduces and validates the effectiveness of K-means clustering as a labeling strategy for web performance categorization. Previous research has predominantly relied on expert-defined thresholds or simple statistical divisions [5, 14]. The demonstration that unsupervised learning can create more effective performance categories advances understanding of how to approach performance classification tasks.

Third, the achievement of 98.47% F1-score significantly exceeds results reported in related research [5, 14, 19, 22], establishing a new benchmark for web performance prediction accuracy. This superior performance validates the methodology and demonstrates the viability of machine learning for highly accurate performance categorization.

Fourth, the comprehensive feature importance analysis provides empirical insights into which performance metrics most strongly predict overall website quality. These findings can inform both future research directions and practical optimization priorities.

**Practical Contributions:**

From a practical standpoint, the developed web platform demonstrates the feasibility of integrating sophisticated machine learning models into user-friendly tools accessible to developers without specialized ML expertise. This democratization of advanced analytics has the potential to significantly improve the accessibility of performance optimization best practices.

The automated nature of the system addresses the time and resource constraints that often prevent thorough performance optimization. By providing instant analysis and recommendations, the tool enables more frequent performance audits and faster iteration on optimization strategies.

The high accuracy of predictions enables developers to make data-driven decisions about performance optimization priorities with confidence. Rather than relying solely on general best practices, developers can receive specific, quantified predictions about their website's performance category.

**Industry Impact:**

For the web development industry, this research provides a foundation for the next generation of automated performance optimization tools. The demonstrated effectiveness of ML-based prediction could be integrated into continuous integration/continuous deployment (CI/CD) pipelines, enabling automated performance validation before code deployment.

The economic implications are substantial. Given that performance directly impacts conversion rates and user retention [2, 7], tools that make optimization more accessible and effective can have measurable business impact. The automation of analysis reduces the need for expensive manual audits while potentially achieving better results.

**Long-term Implications:**

This research contributes to the broader trend toward intelligent, automated web development tools. As websites continue to increase in complexity, automated analysis and optimization will become increasingly necessary. The methodologies and findings from this research provide a template for applying machine learning to other aspects of web development and quality assurance.

Furthermore, the research demonstrates that domain-specific applications of machine learning (in this case, web performance) can achieve superior results compared to generic approaches, encouraging further investigation into specialized ML applications in web technologies."""

add_paragraph(doc, intro_1_5)

doc.add_paragraph()

add_heading(doc, "1.6 Scope and Limitations", level=2)

intro_1_6 = """While this research makes significant contributions, it is important to acknowledge its scope and limitations:

**Scope:**

This research focuses specifically on client-side performance metrics as measured through tools compatible with Google's Lighthouse API [17]. The study encompasses:
- 1,167 websites across diverse domains and industries
- 22 distinct performance features including all Core Web Vitals
- Three labeling strategies and three ML algorithms (nine total models)
- Desktop and mobile web performance (not native mobile applications)
- Static snapshot measurements rather than continuous monitoring

**Limitations:**

1. **Dataset Constraints:** While comprehensive, the dataset of 1,167 websites, though substantial, may not capture all possible variations in web technologies, frameworks, and architectures present across the billions of websites globally. The dataset represents a snapshot in time and does not account for temporal variations in performance.

2. **Geographic Considerations:** Performance measurements were conducted from specific geographic locations and network conditions. Actual user experience may vary significantly based on geographic proximity to servers, local network conditions, and Content Delivery Network (CDN) configurations.

3. **Technology Evolution:** Web technologies evolve rapidly. Models trained on current data may require retraining as new frameworks, browser capabilities, and optimization techniques emerge. The model's applicability to future web technologies is subject to this limitation.

4. **Categorical Prediction:** The current models predict performance categories (Good, Average, Weak) rather than exact metric values. While highly accurate for classification, they do not provide precise predictions of specific metrics like exact LCP times.

5. **Context Independence:** The models do not account for business-specific context such as target audience characteristics, acceptable performance thresholds for specific industries, or the relative importance of different metrics for particular business models.

6. **Implementation Complexity:** While the web platform demonstrates practical applicability, full production deployment at scale would require additional considerations for load balancing, caching, model versioning, and monitoring that are beyond the scope of this research.

7. **Validation Scope:** Real-world validation was conducted on a limited scale. Large-scale deployment with thousands of daily users would provide additional insights into model robustness and practical utility.

These limitations provide opportunities for future research and development, as discussed in Chapter 6."""

add_paragraph(doc, intro_1_6)

doc.add_paragraph()

add_heading(doc, "1.7 Thesis Organization", level=2)

intro_1_7 = """This thesis is organized into six chapters, structured to provide a comprehensive account of the research from motivation through conclusions:

**Chapter 1: Introduction** provides the background and motivation for the research, establishes the problem statement, outlines research objectives and questions, discusses the significance of the study, and defines its scope and limitations.

**Chapter 2: Literature Review** presents a comprehensive review of relevant literature including web performance metrics and Core Web Vitals [8, 9, 20], machine learning applications in web technologies [5, 13, 16], related research on performance prediction [14, 19, 22], and identifies gaps in existing research that this study addresses.

**Chapter 3: Research Methodology** details the research design, describes the dataset and data collection procedures, explains the three labeling strategies (tertile-based, weighted composite, K-means clustering), presents the three machine learning algorithms employed (Random Forest [4], LightGBM [13], Neural Networks [16]), and describes the evaluation methodology.

**Chapter 4: System Design and Implementation** presents the overall system architecture, details the data processing pipeline, describes the model training implementation, explains the web platform development including frontend (Next.js) and backend (FastAPI) components, and discusses deployment considerations.

**Chapter 5: Results and Discussion** presents the dataset characteristics, provides comprehensive performance comparison of all nine models, analyzes confusion matrices and feature importance, compares results with existing research [5, 14, 19], presents real-world application results, and discusses the implications of findings.

**Chapter 6: Conclusion and Future Work** summarizes the research findings, articulates the specific contributions made, acknowledges limitations, provides recommendations for future research directions, and offers concluding remarks on the significance and impact of this work.

The thesis concludes with a comprehensive **References** section containing all cited works, followed by **Appendices** that provide detailed code implementations and additional visualizations."""

add_paragraph(doc, intro_1_7)

doc.add_page_break()

print("📄 Creating Chapter 2: Literature Review...")

# ========================================
# CHAPTER 2: LITERATURE REVIEW
# ========================================

add_heading(doc, "CHAPTER 2", level=1)
add_heading(doc, "LITERATURE REVIEW", level=1)

add_heading(doc, "2.1 Web Performance Metrics", level=2)

lit_2_1 = """Web performance measurement has evolved significantly over the past two decades, reflecting growing understanding of what constitutes meaningful user experience. Early performance metrics focused primarily on technical measurements such as page load time and bandwidth utilization [3]. While these metrics provided useful technical insights, they often failed to correlate strongly with actual user experience, leading to the development of more sophisticated, user-centric metrics [19].

Traditional performance metrics include:

**Page Load Time:** The total time from initiating a page request to complete page load. While intuitive, this metric has limitations as it doesn't capture when content becomes visible or interactive to users [20].

**Time to First Byte (TTFB):** Measures server response time—the duration from request initiation to receiving the first byte of data. This metric primarily reflects server performance and network latency but doesn't indicate when users can actually interact with content [20].

**DOMContentLoaded:** Indicates when the initial HTML document has been completely loaded and parsed. However, this doesn't necessarily correspond to visual completeness or interactivity, as stylesheets, images, and scripts may still be loading.

**Load Event:** Fires when all resources (images, scripts, stylesheets) have loaded. This metric often occurs well after the page appears complete to users, making it a poor indicator of perceived performance.

The limitations of these traditional metrics led to the development of user-centric performance metrics that better reflect actual user experience. Google has been at the forefront of this evolution, introducing metrics such as First Contentful Paint (FCP), which measures when the first content element appears on screen, providing a better indication of when users first see feedback from their page request [9].

Speed Index, developed as part of the WebPageTest project, measures how quickly content is visually displayed during page load. It provides a more nuanced view of loading performance by considering the progression of visual completeness rather than just final load state [19].

The Navigation Timing API [21] and subsequent Performance APIs standardized by the W3C have enabled precise measurement of various performance milestones, providing the technical foundation for modern performance monitoring tools like Google Lighthouse [17]."""

add_paragraph(doc, lit_2_1)

doc.add_paragraph()

add_heading(doc, "2.2 Core Web Vitals and User Experience", level=2)

lit_2_2 = """In May 2020, Google introduced Core Web Vitals as a set of standardized metrics designed to capture essential aspects of user experience [8]. This initiative marked a significant evolution in web performance measurement by focusing on metrics that directly correlate with user experience quality and making these metrics ranking factors in Google's search algorithm [10].

Core Web Vitals consist of three key metrics:

**Largest Contentful Paint (LCP)** measures loading performance by identifying when the largest content element in the viewport becomes visible [9]. Good LCP is defined as 2.5 seconds or less from page load initiation. LCP represents a significant improvement over traditional load metrics because it measures when the main content becomes visible to users, which typically correlates with their perception of loading speed. Research by Google has demonstrated strong correlation between LCP and user engagement metrics [7].

**First Input Delay (FID) / Interaction to Next Paint (INP)** measures interactivity by quantifying the delay between a user's first interaction with a page and the browser's response to that interaction [9]. As of March 2024, FID has been replaced by INP, which provides a more comprehensive measure of overall page responsiveness throughout the entire page lifecycle. Good INP is defined as 200 milliseconds or less. This metric addresses the critical user experience issue of pages that appear loaded but remain unresponsive to user input due to JavaScript execution or other main thread blocking activities.

**Cumulative Layout Shift (CLS)** measures visual stability by quantifying unexpected layout shifts that occur during page load and throughout the page lifecycle [9]. Good CLS is defined as 0.1 or less. Unexpected layout shifts can cause users to accidentally click on wrong elements, particularly problematic on mobile devices. CLS has become increasingly important as websites incorporate dynamic content, advertisements, and third-party widgets that can cause layout instability.

The introduction of Core Web Vitals as ranking factors represents Google's recognition that page experience significantly impacts user satisfaction and engagement. Research by Google demonstrated that sites meeting Core Web Vitals thresholds have 24% lower abandonment rates [7]. This business impact has made Core Web Vitals a priority for website owners and developers.

Beyond the three core metrics, Google's Web Vitals initiative includes additional metrics such as First Contentful Paint (FCP), measuring when any content first appears on screen, and Time to Interactive (TTI), measuring when a page becomes fully interactive [9]. While not designated as "core" metrics, these additional measurements provide valuable insights into different aspects of performance and user experience.

The standardization of Core Web Vitals has facilitated research on web performance by providing consistent, well-defined metrics that can be measured across different websites and contexts. Tools like Google Lighthouse [17] and the Chrome User Experience Report provide standardized measurement methodologies, enabling comparative research such as the present study."""

add_paragraph(doc, lit_2_2)

doc.add_page_break()

# Continue with Chapter 3 and add images
print("📄 Creating Chapter 3: Research Methodology...")

add_heading(doc, "CHAPTER 3", level=1)
add_heading(doc, "RESEARCH METHODOLOGY", level=1)

add_heading(doc, "3.1 Research Design and Approach", level=2)

meth_3_1 = """This research employs a quantitative, experimental methodology to develop and evaluate machine learning models for web performance prediction. The research design follows a systematic pipeline approach, ensuring reproducibility and scientific rigor while enabling comprehensive comparison of different methodological choices.

The overall research approach can be characterized as follows:

**Paradigm:** Positivist, quantitative research utilizing empirical data and statistical validation.

**Type:** Applied research with both theoretical contributions (comparative analysis of algorithms and labeling strategies) and practical outputs (functional web platform).

**Strategy:** Experimental design with controlled comparison of multiple machine learning approaches.

**Method:** Supervised learning for classification, supplemented by unsupervised learning (K-means) for one labeling strategy.

The research methodology consists of six main phases:

1. **Data Collection and Dataset Preparation**: Gathering performance metrics from 1,167 diverse websites using automated tools compatible with Google Lighthouse API [17].

2. **Data Preprocessing and Exploration**: Cleaning the dataset, handling missing values, detecting outliers, and conducting exploratory data analysis to understand feature distributions and relationships.

3. **Labeling Strategy Development**: Implementing three distinct approaches to categorize website performance: tertile-based division, weighted composite scoring, and K-means clustering.

4. **Feature Engineering and Selection**: Processing raw metrics, creating derived features where appropriate, and selecting the final feature set for model training.

5. **Model Development and Training**: Training nine distinct models (three labeling strategies × three algorithms) using standardized training procedures to enable fair comparison.

6. **Evaluation and Validation**: Assessing model performance using multiple metrics (accuracy, precision, recall, F1-score) and validating results through confusion matrix analysis and real-world testing.

This structured approach enables systematic comparison while maintaining scientific rigor. Each phase builds upon the previous one, with careful documentation of decisions and parameters to ensure reproducibility. The use of fixed random seeds (random_state=42) throughout the experiment ensures that results can be replicated by other researchers."""

add_paragraph(doc, meth_3_1)

doc.add_paragraph()

# Save progress and continue
output_path = "f:/client/Optimizer/optimizer/paperrs/Complete_Thesis_MD_ASHIKUR_RAHMAN.docx"
doc.save(output_path)

print("\n" + "="*80)
print("✅ PART 2A SAVED - Chapters 1-3 Started")
print("="*80)
print(f"File: {output_path}")
print("Continuing with methodology details and results with ALL IMAGES...")
print("="*80)
